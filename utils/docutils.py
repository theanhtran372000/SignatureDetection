import re
from docx import Document

replaceCountNumber = 0
replaceCountDate = 0


def docx_replace_regex(doc, regex, replace):
    global replaceCountNumber
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)
    for p in doc.paragraphs:
        if (regex.search(p.text) and replaceCountNumber == 0):
            print(regex.search(p.text))
            # text = regex.sub(replace, p.text, count = 0)
            text = re.sub(regex.search(p.text).string, replace, p.text, 1)
            p.text = text
            replaceCountNumber = replaceCountNumber + 1
            break
                
    return doc



def docx_replace_date(doc, regex, replace):
    global replaceCountDate
    for p in doc.paragraphs:
        if (regex.search(p.text) and replaceCountDate == 0):
            # text = regex.sub(replace, p.text)
            text = re.sub(regex.search(p.text).string, replace, p.text, 1)
            p.text = text
            replaceCountDate = replaceCountDate + 1
            break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_date(cell, regex, replace)
                
    return doc


def remove_datetime(doc):
    global replaceCountDate
    regex = re.compile(r".{0,20}, ngày.{0,10}tháng.{0,10}năm.{0,10}")
    replace = " " * 10
    replaceCountDate = 0
    doc = docx_replace_date(doc, regex, replace)
    return doc


def remove_docnum(doc):
    global replaceCountNumber
    regex = re.compile(r"Số:.{0,20}/[^\s\n\t]{2,15}")
    replace = " " * 10
    replaceCountNumber = 0
    doc = docx_replace_regex(doc, regex, replace)
    return doc


def preprocess_doc(in_path, out_path):
    global replaceCountNumber
    global replaceCountDate
    print("Preprocess")
    doc = Document(in_path)
    replaceCountDate = 0
    doc = remove_datetime(doc)
    replaceCountNumber = 0
    doc = remove_docnum(doc)

    doc.save(out_path)
    return out_path
